#!/usr/bin/env python3
"""Convert markdown to DOCX with stable native table rendering."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Cm


def is_table_separator(line: str) -> bool:
    s = line.strip()
    if not s.startswith("|"):
        return False
    parts = [p.strip() for p in s.strip("|").split("|")]
    if not parts:
        return False
    return all(re.fullmatch(r":?-{3,}:?", p or "") for p in parts)


def parse_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def set_landscape(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width


def apply_qc_table_style(table) -> None:
    table.style = "Table Grid"
    table.autofit = False
    widths = [2.2, 2.2, 2.4, 5.5, 4.0, 4.0, 1.8, 2.5]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths):
                cell.width = Cm(widths[idx])


def markdown_to_docx(src: Path, dest: Path) -> None:
    lines = src.read_text(encoding="utf-8").splitlines()

    doc = Document()
    set_landscape(doc.sections[0])

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line.strip():
            doc.add_paragraph("")
            i += 1
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue

        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        if line.strip().startswith("- [ ] "):
            doc.add_paragraph("□ " + line.strip()[6:].strip())
            i += 1
            continue

        if line.strip().startswith("- "):
            doc.add_paragraph(line.strip()[2:].strip(), style="List Bullet")
            i += 1
            continue

        # Markdown table block: header + separator + rows
        if "|" in line and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            header = parse_table_row(line)
            rows: list[list[str]] = []
            i += 2
            while i < len(lines):
                current = lines[i].strip()
                if not current.startswith("|"):
                    break
                rows.append(parse_table_row(lines[i]))
                i += 1

            table = doc.add_table(rows=1 + len(rows), cols=len(header))
            apply_qc_table_style(table)

            for col_idx, value in enumerate(header):
                run = table.cell(0, col_idx).paragraphs[0].add_run(value)
                run.bold = True

            for row_idx, row in enumerate(rows, start=1):
                for col_idx, value in enumerate(row):
                    if col_idx < len(header):
                        table.cell(row_idx, col_idx).text = value

            continue

        doc.add_paragraph(line)
        i += 1

    dest.parent.mkdir(parents=True, exist_ok=True)
    doc.save(dest)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Convert markdown to table-safe DOCX")
    parser.add_argument("source", type=Path, help="Source markdown path")
    parser.add_argument("output", type=Path, nargs="?", help="Output DOCX path")
    args = parser.parse_args()

    src = args.source.resolve()
    out = args.output.resolve() if args.output else src.with_suffix(".docx")
    markdown_to_docx(src, out)
